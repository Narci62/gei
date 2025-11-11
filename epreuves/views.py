import os
import shutil
import subprocess
import traceback
from uuid import uuid4
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.core.files import File
from .forms import EpreuveForm
from .models import Epreuve
from pedagogie.models import Exercice
from comptes.decorators import abonnement_requis
import chardet


@login_required
@abonnement_requis
def generer_epreuve(request):
    matiere = request.POST.get("matiere")
    classe = request.POST.get("classe")

    form = EpreuveForm(request.POST or None, matiere=matiere, classe=classe)

    if request.method == "POST" and form.is_valid():
        professeur = request.user
        matiere = form.cleaned_data["matiere"]
        classe = form.cleaned_data["classe"]
        trimestre = form.cleaned_data["trimestre"]
        themes = form.cleaned_data["themes"]
        lecons = form.cleaned_data["lecons"]

        # Sélection aléatoire des exercices
        exo1 = Exercice.objects.filter(
            type_exercice="situation_probleme", matiere=matiere, classe=classe,
            theme__in=themes, lecon__in=lecons, trimestre=trimestre
        ).order_by('?').first()

        partie_a = Exercice.objects.filter(
            type_exercice="partie_a", matiere=matiere, classe=classe,
            theme__in=themes, lecon__in=lecons, trimestre=trimestre
        ).order_by('?').first()

        partie_b = Exercice.objects.filter(
            type_exercice="partie_b", matiere=matiere, classe=classe,
            theme__in=themes, lecon__in=lecons, trimestre=trimestre
        ).order_by('?').first()

        partie_c = Exercice.objects.filter(
            type_exercice="partie_c", matiere=matiere, classe=classe,
            theme__in=themes, lecon__in=lecons, trimestre=trimestre
        ).order_by('?').first()

        exo3 = Exercice.objects.filter(
            type_exercice="traditionnelle", matiere=matiere, classe=classe,
            theme__in=themes, lecon__in=lecons, trimestre=trimestre
        ).order_by('?').first()

        exercices = [exo1, partie_a, partie_b, partie_c, exo3]

        if not all(exercices):
            return render(request, "generer_epreuve.html", {
                "form": form,
                "error_message": "Pas assez d'exercices disponibles."
            })

        # Vérifier l'existence des fichiers .tex
        for i, exo in enumerate(exercices, start=1):
            if not exo.fichier_tex or not os.path.exists(exo.fichier_tex.path):
                return render(request, "generer_epreuve.html", {
                    "form": form,
                    "error_message": f"Le fichier .tex de l'exercice {i} est manquant ou invalide."
                })

        uid = str(uuid4())
        temp_dir = os.path.join(settings.BASE_DIR, "temp_latex", uid)
        os.makedirs(temp_dir, exist_ok=True)

        noms_fichiers = []
        # Copier et convertir chaque fichier .tex en UTF-8
        for i, exo in enumerate(exercices, start=1):
            nom = f"exercice{i}.tex"
            dst = os.path.join(temp_dir, nom)

            # Lecture binaire pour détection d'encodage
            with open(exo.fichier_tex.path, "rb") as src_file:
                raw = src_file.read()
                detected_enc = chardet.detect(raw)["encoding"] or "latin-1"
                text = raw.decode(detected_enc)

            # Réécriture en UTF-8
            with open(dst, "w", encoding="utf-8") as dst_file:
                dst_file.write(text)

            noms_fichiers.append(nom)

        # Création du fichier LaTeX principal
        main_tex_path = os.path.join(temp_dir, "epreuve.tex")
        with open(main_tex_path, "w", encoding="utf-8") as f:
            f.write(r"""\documentclass[11pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[french]{babel}
\usepackage[T1]{fontenc}
\usepackage{amsmath,amssymb,amsfonts}
\usepackage{graphicx,multicol}
\usepackage{pgfplots}
\pgfplotsset{compat=1.15}
\usepackage{mathrsfs}
\usetikzlibrary{arrows,patterns}
\usepackage{fourier}
\mathversion{bold}
\everymath{\displaystyle}
\usepackage[left=2cm,right=1.5cm,top=1cm,bottom=1.5cm]{geometry}
\usepackage{fancyhdr}
\pagestyle{fancy}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}
\lhead{}\chead{}\rhead{}
\lfoot{}\cfoot{}\rfoot{}
\begin{document}
\noindent\begin{tabular}{|p{5cm}|p{7cm}|p{4cm}|}
\hline 
\textbf{IESG:} & \textbf{Devoir Surveillé} &  \textbf{Classe : """ + classe.nom + r"""}\\ 
\hline 
\textbf{AN/SC : } &\textbf{Épreuve de """ + matiere.nom + r"""} &  \textbf{ Durée :}\\ 
\hline 
\end{tabular} \\

""")
            f.write(r"\noindent\textbf{\underline{EXERCICE 1}}  \textit{(8pts)}\\" + "\n")
            f.write(r"\input{exercice1.tex}" + "\n\n")
            f.write(r"\noindent\textbf{\underline{EXERCICE 2}}  \textit{(6pts)}\\" + "\n")
            f.write(r"\noindent\textbf{\underline{Partie A}}  \textit{(2pts)}\\" + "\n")
            f.write(r"\input{exercice2.tex}" + "\n\n")
            f.write(r"\noindent\textbf{\underline{Partie B}}  \textit{(2pts)}\\" + "\n")
            f.write(r"\input{exercice3.tex}" + "\n\n")
            f.write(r"\noindent\textbf{\underline{Partie C}}  \textit{(2pts)}\\" + "\n")
            f.write(r"\input{exercice4.tex}" + "\n\n")

            f.write(r"\noindent\textbf{\underline{EXERCICE 3}}  \textit{(6pts)}\\" + "\n")
            f.write(r"\input{exercice5.tex}" + "\n\n")
            f.write(r"\end{document}")

        try:
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-output-directory", temp_dir, main_tex_path],
                capture_output=True,
                text=True
            )

            pdf_path = os.path.join(temp_dir, "epreuve.pdf")
            if os.path.exists(pdf_path):
                # Sauvegarde dans la base
                epreuve = Epreuve.objects.create(
                    professeur=professeur,
                    matiere=matiere,
                    classe=classe,
                    trimestre=trimestre.nom
                )
                with open(pdf_path, "rb") as f:
                    epreuve.fichier_pdf.save(f"epreuve_{uid}.pdf", File(f), save=True)

                shutil.rmtree(temp_dir)
                return redirect("mes_epreuves")

            return HttpResponse("Erreur LaTeX : PDF non généré.")

        except Exception as e:
            traceback.print_exc()
            return HttpResponse("Erreur interne lors de la génération LaTeX.")

    return render(request, "generer_epreuve.html", {"form": form})




from django.core.paginator import Paginator
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import Epreuve
from pedagogie.models import Matiere

@login_required
@abonnement_requis
def mes_epreuves(request):
    professeur = request.user.id
    epreuve_list = Epreuve.objects.filter(professeur=professeur).order_by('-date_creation')

    # Récupération des filtres depuis GET
    matiere_id = request.GET.get('matiere')
    trimestre_val = request.GET.get('trimestre')

    # Application des filtres si présents
    if matiere_id:
        epreuve_list = epreuve_list.filter(matiere_id=matiere_id)

    if trimestre_val:
        epreuve_list = epreuve_list.filter(trimestre=trimestre_val)

    # Pagination : 5 épreuves par page
    paginator = Paginator(epreuve_list, 5)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Pour le formulaire de filtre
    matieres = Matiere.objects.all()
    trimestres = Trimestre.objects.values_list('nom', flat=True).distinct()

    context = {
        "page_obj": page_obj,
        "matieres": matieres,
        "trimestres": trimestres,
    }

    return render(request, "mes_epreuves.html", context)







from django.http import JsonResponse
from pedagogie.models import Lecon, Trimestre
from .models import Theme

def get_lecons_by_themes(request):
    theme_ids = request.GET.getlist('theme_ids[]')
    trimestre_id = request.GET.get('trimestre')

    if not theme_ids:
        return JsonResponse([], safe=False)

    lecons = Lecon.objects.filter(theme__id__in=theme_ids)

    if trimestre_id:
        try:
            trimestre_obj = Trimestre.objects.get(id=trimestre_id)
            numero_trimestre = trimestre_obj.numero  # Assure-toi que le modèle a un champ `numero`
            lecons = lecons.filter(trimestre__numero__lte=numero_trimestre)
        except Trimestre.DoesNotExist:
            pass

    data = [{"id": l.id, "nom": l.nom} for l in lecons.distinct()]
    return JsonResponse(data, safe=False)

from docx import Document
from django.http import HttpResponse
from .models import Epreuve

def exporter_epreuve_word(request, epreuve_id):
    epreuve = Epreuve.objects.get(id=epreuve_id)

    document = Document()
    document.add_heading(f"Epreuve : {epreuve.code}", 0)

    document.add_paragraph(f"Matière : {epreuve.matiere.nom}")
    document.add_paragraph(f"Classe : {epreuve.classe.nom}")
    document.add_paragraph(f"Trimestre : {epreuve.trimestre}")
    document.add_paragraph("")

    document.add_heading("Exercices :", level=1)
    for exercice in epreuve.exercices.all():
        document.add_paragraph(f"{exercice.intitule}", style='List Bullet')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=epreuve_{epreuve.code}.docx'
    document.save(response)
    return response

